"""Build report/report.docx — the HyperInvest Part B report scaffold.

Structure follows PROJECT_BRIEF.md Section 5's suggested six sections, and
every section is annotated with the HD-band cell of its marking-rubric
criterion (Section 9) plus the real, verified numbers from results/ (read
from the committed CSVs at generation time, never hand-transcribed).

The bracketed [RUBRIC], [EVIDENCE] and [WRITE] blocks are guidance for the
student, who rewrites every section in his own words and deletes the
guidance as he goes. Run from the project root:

    python scripts/build_report_scaffold.py
"""
from __future__ import annotations

import pathlib

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
RES = ROOT / "results"
FIG = RES / "figures"
OUT = ROOT / "report" / "report.docx"

metrics = pd.read_csv(RES / "tables" / "performance_metrics.csv", index_col="fund")
fusion = pd.read_csv(RES / "tables" / "fusion_comparison.csv", index_col="metric")
compare = pd.read_csv(RES / "tables" / "sentiment_model_comparison.csv", index_col="sector")
sent_sum = pd.read_csv(RES / "tables" / "sector_sentiment_summary.csv", index_col=0)

GREY = RGBColor(0x6B, 0x62, 0x5C)
TEAL = RGBColor(0x0F, 0x76, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def rubric(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("[RUBRIC — what HD requires] " + text)
    r.italic = True
    r.font.color.rgb = TEAL
    r.font.size = Pt(9.5)


def evidence(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("[EVIDENCE — real numbers from results/] " + text)
    r.font.color.rgb = GREY
    r.font.size = Pt(9.5)


def write_note(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("[WRITE] " + text)
    r.bold = True
    r.font.size = Pt(9.5)


def exhibit(path: pathlib.Path, caption: str, width: float = 6.0) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.add_paragraph(f"Exhibit: {path.name}. {caption}").style = doc.styles[
            "Normal"]
    else:
        doc.add_paragraph(f"[MISSING EXHIBIT: {path.name}]")


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


# ---------------------------------------------------------------------------
# Title
# ---------------------------------------------------------------------------
doc.add_heading("HyperInvest — Systematic Funds, News Sentiment, and a "
                "Learning Sandbox", level=0)
doc.add_paragraph("FINS3645 Project B · z5538441 · 2026")
doc.add_paragraph(
    "Guidance blocks in [brackets] are scaffolding: they state the rubric "
    "target and the verified evidence for each section. Replace them with "
    "your own analysis and delete every bracketed block before submission. "
    "Limit: max 10 pages / ~5,000 words excluding appendix and references. "
    "Every exhibit must be referenced and interpreted in the text.")
doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. The funds and the backtest design
# ---------------------------------------------------------------------------
doc.add_heading("1. The funds and the backtest design", level=1)
rubric("Funds criterion (15%), HD: equity-only, crypto-only AND combined "
       "funds across several optimisation methods, each with a correct "
       "walk-forward out-of-sample backtest (no look-ahead, weights from "
       "past data only, correct 252 vs 365 annualisation).")
evidence(
    "Eight funds: Combined (60 assets) x {Equal-Weight, Min-Variance, "
    "Max-Sharpe, Risk-Parity}; Equity (50) x {Max-Sharpe, Min-Variance}; "
    "Crypto (10) x {Max-Sharpe}. Walk-forward: 252-trading-day estimation "
    "window; monthly rebalance on the last trading day; weights use only "
    "the trailing window (rebalance day excluded); long-only; risk-free "
    "rate 0 and zero transaction costs (both stated assumptions). Max-Sharpe "
    "is solved as the convex QP tangency reformulation after the direct "
    "ratio optimisation stalled (Section 6 material). Annualisation: 252 "
    "for equity-calendar funds, 365 only for Crypto Max-Sharpe, which runs "
    "on its own 365-day calendar (OOS starts 2020-10-01 vs 2021-02-01 for "
    "the others).")
write_note("Motivate the fund shelf for a novice investor, then state the "
           "backtest rules precisely — window, rebalance rule, constraints, "
           "assumptions — and why each rule prevents look-ahead. Name the "
           "week-4/5 course concepts you applied (estimation windows, "
           "formation/return-date discipline) and flag Ledoit-Wolf-style "
           "shrinkage and the QP reformulation as beyond-course choices you "
           "made deliberately, with the reason.")

# ---------------------------------------------------------------------------
# 2. Out-of-sample results and fund fact sheets
# ---------------------------------------------------------------------------
doc.add_heading("2. Out-of-sample results and fund fact sheets", level=1)
rubric("Same criterion, HD continued: fact sheets and the required exhibits "
       "(metrics table, growth of $1, drawdown, weights over time) present, "
       "and the funds are COMPARED — not just listed.")
evidence(
    f"performance_metrics.csv: best Sharpe = Combined Risk-Parity "
    f"{metrics.loc['Combined Risk-Parity','sharpe']:.2f} "
    f"(ann. return {pct(metrics.loc['Combined Risk-Parity','ann_return'])}, "
    f"vol {pct(metrics.loc['Combined Risk-Parity','ann_vol'])}, max DD "
    f"{pct(metrics.loc['Combined Risk-Parity','max_drawdown'])}); Combined "
    f"Max-Sharpe Sharpe {metrics.loc['Combined Max-Sharpe','sharpe']:.2f} "
    f"with the highest combined-fund return "
    f"{pct(metrics.loc['Combined Max-Sharpe','ann_return'])}; Combined "
    f"Min-Variance has the smallest drawdown "
    f"({pct(metrics.loc['Combined Min-Variance','max_drawdown'])}) and "
    f"lowest vol ({pct(metrics.loc['Combined Min-Variance','ann_vol'])}). "
    f"Crypto Max-Sharpe: {pct(metrics.loc['Crypto Max-Sharpe','ann_return'])} "
    f"return with {pct(metrics.loc['Crypto Max-Sharpe','ann_vol'])} vol and "
    f"{pct(metrics.loc['Crypto Max-Sharpe','max_drawdown'])} max DD. "
    "All equity-calendar funds: 734 OOS trading days, 2021-02-01 to "
    "2023-12-29; crypto fund 1,187 days.")
doc.add_heading("Exhibit 1 — Performance metrics across funds (table)", level=3)
evidence("Insert results/tables/performance_metrics.csv formatted as a "
         "Word table here (all 8 rows).")
exhibit(FIG / "growth_of_1_all_funds.png",
        "Exhibit 2 — Growth of $1, all eight funds, OOS period.")
exhibit(FIG / "sharpe_by_fund.png",
        "Exhibit 3 — Sharpe ratio by fund.")
write_note("Compare, don't list: which method family behaved how, and WHY "
           "economically (e.g. why min-variance protected capital in 2022; "
           "why max-Sharpe concentrated; why crypto's calendar and vol "
           "dominate). Interpret every exhibit. Use week-3 evaluation "
           "discipline (compare against benchmarks, not absolute numbers).")

# ---------------------------------------------------------------------------
# 3. The sentiment index
# ---------------------------------------------------------------------------
doc.add_heading("3. The news-sentiment index", level=1)
rubric("Sentiment criterion (10%), HD: a sentiment model applied to the "
       "headlines building a VALIDATED standalone sector index shown over "
       "time; look-ahead-safe handling justified.")
evidence(
    "VADER + custom ~60-term finance lexicon (FINANCE_LEXICON), scores on "
    "VADER's -4..+4 scale; ticker-day average, 5-day EMA, equal-weight "
    "sector average, forward-fill <=10 trading days then neutral, lagged 1 "
    "trading day (so day t uses only information to t-1). Index: 2020-01-02 "
    "to 2023-12-29 (1,006 days). Positive level bias measured: sector means "
    f"+{sent_sum['mean'].min():.2f} to +{sent_sum['mean'].max():.2f}, only "
    f"{sent_sum['pct_below_zero'].min():.1f}-"
    f"{sent_sum['pct_below_zero'].max():.1f}% of readings below zero - the "
    "app therefore offers the week-9 standardised (z-score) view. "
    "VALIDATION against the course's finVADER benchmark on the identical "
    f"pipeline: sign agreement {compare.loc['ALL','sign_agreement_pct']:.1f}%, "
    f"per-sector Pearson r {compare['pearson_r'].min():.2f}-"
    f"{compare['pearson_r'].max():.2f} (all-sector r "
    f"{compare.loc['ALL','pearson_r']:.2f}), 1,006 overlapping days.")
exhibit(FIG / "sector_sentiment_index.png",
        "Exhibit 4 — Sector news-sentiment index over time.")
write_note("Justify each text-handling choice (why headlines are scored "
           "whole, why EMA, why the ffill cap, why the lag). Present the "
           "positive bias and the standardisation fix as something you "
           "measured and handled (week 9). Present the finVADER comparison "
           "as your validation evidence, including one honest weakness "
           "(e.g. the 'despite' negation quirk - see "
           "results/tables/sentiment_disagreement_examples.csv). Cite week "
           "7/8 discipline: coverage analysis as the reason the index is "
           "sector-level; the human-approval gate as the lexicon's "
           "governance model.")

# ---------------------------------------------------------------------------
# 4. Extensions and innovations
# ---------------------------------------------------------------------------
doc.add_heading("4. Extensions and innovations", level=1)
rubric("Innovation criterion (30%, the heaviest), HD: a distinctive, "
       "IMPLEMENTED extension shown with evidence — e.g. a custom sentiment "
       "tool or lexicon, an original evaluation method, a custom design "
       "system, or a genuinely valuable app feature. Built and "
       "demonstrated, not proposed. A careful extension with a negative "
       "result, explained, still earns this band.")
evidence(
    "Four implemented extensions, each with evidence in results/: "
    "(1) the custom finance lexicon, validated against finVADER (Section 3 "
    "numbers); (2) the Practice-layer 'investment time machine': a Replay "
    "sandbox with a draggable timeline over the OOS period, event-driven "
    "time granularity (7 turbulent windows from a 2-sigma rule + named "
    "events), dollar-tracked holdings with buy-and-hold drift, and "
    "just-in-time neutral learning cards; (3) an original design system "
    "(typography, surface palette, one chart language) - the rubric's own "
    "words are 'colour, type, and figure language'; (4) the sentiment-tilt "
    "fusion as a deliberately honest negative result: Sharpe "
    f"{fusion.loc['sharpe','base_equity_max_sharpe']:.3f} base vs "
    f"{fusion.loc['sharpe','equity_sentiment_tilt']:.3f} tilted, ann. "
    f"return {pct(fusion.loc['ann_return','base_equity_max_sharpe'])} vs "
    f"{pct(fusion.loc['ann_return','equity_sentiment_tilt'])}.")
exhibit(FIG / "fusion_growth_of_1.png",
        "Exhibit 5 — Fusion before vs after: growth of $1, base vs tilted.")
write_note("This section carries 30% — make the innovation CASE here. For "
           "each extension: what it is, why it is genuinely valuable, and "
           "the evidence it works (or, for the fusion, what the negative "
           "result teaches: headline tone is a noisy signal, and a naive "
           "tilt can hurt - the brief itself says so). The time machine is "
           "the headline innovation: describe the design decisions (slider "
           "as primary control, event windows, drift made visible, "
           "neutrality-preserving learning cards) and what a user learns "
           "that a static chart cannot teach.")

# ---------------------------------------------------------------------------
# 5. The app and the investor journey
# ---------------------------------------------------------------------------
doc.add_heading("5. The app and the investor journey", level=1)
rubric("App criterion (15%), HD: a reliable app, deployed from a public "
       "GitHub repo, supporting the full investor journey (compare funds, "
       "fact sheet, set allocation) plus sentiment analytics, running on a "
       "basic machine; an original design system strengthens this band.")
evidence(
    "App reads only committed results/ artifacts (no optimiser, no VADER "
    "at runtime; nltk absent from the app by design). Live URL and public "
    "repo link: [INSERT AFTER DEPLOYMENT]. Two journey personas used in "
    "design: the novice (plain-English mode, Practice sandbox, learning "
    "cards) and the experienced self-directed investor (professional mode, "
    "comparison table, fact sheets, allocation). Neutrality stance "
    "throughout: evidence layer, not advice layer.")
write_note("Describe the target user and the customer journey in your own "
           "words (the two personas). Include 2-4 app screenshots [capture "
           "after deployment]. State the neutrality design decisions "
           "explicitly (no recommendations, no scores, 'not a buy or sell "
           "signal' disclaimer) and why they are a professional stance, not "
           "a limitation. Name the deployment topology: precomputed "
           "artifacts + lightweight app = fast free-tier loads.")

# ---------------------------------------------------------------------------
# 6. Critical reflection and recommendations
# ---------------------------------------------------------------------------
doc.add_heading("6. Critical reflection and recommendations", level=1)
rubric("Interpretation criterion (10%), HD: evidence-based reflection on "
       "what worked, what did not, and why, with THREE concrete, specific "
       "real-world recommendations. Every exhibit interpreted; your own "
       "words.")
evidence(
    "Candidate honest material: the fusion underperformed (Section 4); the "
    "max-Sharpe SLSQP stall that forced the QP reformulation (caught by the "
    "weights sanity check); the sentiment model's documented weak spots "
    "(negation quirks, promotional-language bias); zero transaction costs "
    "and a zero risk-free rate are stated simplifications; data ends "
    "2023-12-31 so nothing here describes today's market.")
write_note("Choose YOUR three recommendations — they are graded on being "
           "concrete and yours. Candidates discussed and evidenced in our "
           "results: (a) a transaction-cost/turnover model (weights jump "
           "0-90% in some funds - costs would change conclusions); (b) a "
           "nightly data refresh so the evidence stays current (the app "
           "would still never predict); (c) a 5th optimiser such as "
           "mean-CVaR (week 5) for tail-risk-aware allocation. Reflect on "
           "what you would do differently.")

# ---------------------------------------------------------------------------
# AI workflow statement + appendix + references
# ---------------------------------------------------------------------------
doc.add_heading("AI workflow statement", level=1)
rubric("AI Workflow criterion (20%) is assessed from the AI pack "
       "(AGENTS.md + ai/prompt_log.md), not from this report — but state "
       "the workflow briefly and honestly here.")
write_note("One short paragraph in your own words: you directed a single "
           "AI agent (Kimi Code CLI) for planning, implementation and "
           "verification; you approved every material change and judged "
           "results; the full curated log is in ai/prompt_log.md, and your "
           "own instruction file is AGENTS.md. Note the documented "
           "restart (Session 1) and the corrections you required (log "
           "accuracy, neutrality rewrites).")

doc.add_page_break()
doc.add_heading("Appendix — supporting exhibits", level=1)
exhibit(FIG / "drawdown_combined_max_sharpe.png",
        "A1 — Drawdown, Combined Max-Sharpe.")
exhibit(FIG / "weights_combined_max_sharpe.png",
        "A2 — Target weights per rebalance, Combined Max-Sharpe "
        "(top 5; others pooled in grey).")
exhibit(FIG / "sentiment_model_comparison.png",
        "A3 — Custom lexicon vs finVADER benchmark: per-sector correlation "
        "and example sectors.")
evidence("Also reference: results/tables/fusion_comparison.csv (fusion "
         "before/after table), sentiment_model_comparison.csv, "
         "sentiment_disagreement_examples.csv, weights_sanity_check.csv. "
         "Format the fusion table as a Word table either in Section 4 or "
         "here.")

doc.add_heading("References", level=1)
write_note("Cite: the course data bundle; VADER (Hutto & Gilbert 2014); "
           "finVADER / SentiBigNomics / Henry's lexicon as used for "
           "benchmarking; any optimisation references you use. Verify "
           "every reference exists - context/verify_ai_output.md rules.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"scaffold written to {OUT}")
print(f"paragraphs: {len(doc.paragraphs)}")
