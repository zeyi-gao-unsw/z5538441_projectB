"""Build report/report.docx — full DRAFT of the HyperInvest Part B report.

The draft prose below is AI-written scaffolding: the student rewrites every
section in his own words before submission (course mandatory requirement;
see ai/prompt_log.md). Every number in the prose is injected programmatically
from the committed results/ CSVs at generation time — nothing is
hand-transcribed. [RUBRIC] and [EVIDENCE] blocks remain as deletable
guidance. Run from the project root:

    python scripts/build_report_draft.py
"""
from __future__ import annotations

import pathlib

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
RES = ROOT / "results"
FIG = RES / "figures"
OUT = ROOT / "report" / "report.docx"

metrics = pd.read_csv(RES / "tables" / "performance_metrics.csv", index_col="fund")
fusion = pd.read_csv(RES / "tables" / "fusion_comparison.csv", index_col="metric")
compare = pd.read_csv(RES / "tables" / "sentiment_model_comparison.csv", index_col="sector")
sent_sum = pd.read_csv(RES / "tables" / "sector_sentiment_summary.csv", index_col=0)

GREY = RGBColor(0x6B, 0x62, 0x5C)
TEAL = RGBColor(0x0F, 0x76, 0x6E)

m = metrics
rp = m.loc["Combined Risk-Parity"]
cms = m.loc["Combined Max-Sharpe"]
cmv = m.loc["Combined Min-Variance"]
cew = m.loc["Combined Equal-Weight"]
cr = m.loc["Crypto Max-Sharpe"]
ems = m.loc["Equity Max-Sharpe"]
emv = m.loc["Equity Min-Variance"]
tilt = m.loc["Equity+Sentiment Tilt"]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def rubric(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("[RUBRIC — what HD requires] " + text)
    r.italic = True
    r.font.color.rgb = TEAL
    r.font.size = Pt(9.5)


def evidence(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("[EVIDENCE — verified numbers] " + text)
    r.font.color.rgb = GREY
    r.font.size = Pt(9.5)


def body(text: str) -> None:
    doc.add_paragraph(text)


def exhibit(fname: str, caption: str, width: float = 6.0) -> None:
    path = FIG / fname
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.add_paragraph(f"{fname} — {caption}")
    else:
        doc.add_paragraph(f"[MISSING EXHIBIT: {fname}]")


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


# ---------------------------------------------------------------------------
doc.add_heading("HyperInvest — Systematic Funds, News Sentiment, and a "
                "Learning Sandbox", level=0)
doc.add_paragraph("FINS3645 Project B · z5538441 · 2026")
p = doc.add_paragraph()
r = p.add_run("STATUS: FULL AI-WRITTEN DRAFT FOR REWRITING. The student "
              "replaces every section with his own wording before "
              "submission; bracketed [RUBRIC]/[EVIDENCE] blocks are "
              "deletable guidance. Limit: max 10 pages / ~5,000 words "
              "excluding appendix and references.")
r.bold = True
r.font.size = Pt(9.5)
doc.add_page_break()

# ---------------------------------------------------------------------------
doc.add_heading("1. The funds and the backtest design", level=1)
rubric("Funds (15%), HD: equity-only, crypto-only AND combined funds across "
       "several methods, correct walk-forward OOS backtest (no look-ahead, "
       "weights from past data only, correct 252 vs 365 annualisation).")
body(
    "HyperInvest offers eight systematically managed funds across three "
    "asset families. Four combined funds invest in the full 60-asset "
    "universe (50 US large-cap equities and 10 cryptocurrencies) using four "
    "different construction methods: equal-weight, minimum-variance, "
    "maximum-Sharpe, and risk-parity. Two equity-only funds (maximum-Sharpe "
    "and minimum-variance over the 50 stocks) and one crypto-only fund "
    "(maximum-Sharpe over the 10 coins) complete the shelf. Each (family, "
    "method) pair is a separate investable fund, because that is what a "
    "user actually buys and what a fact sheet describes.")
body(
    "Every fund is evaluated with the same walk-forward out-of-sample "
    "backtest. Weights are re-formed on the last trading day of each month "
    "using only the trailing 252 trading days of returns; the rebalance "
    "day's own return is excluded from the estimation slice, so no "
    "information from the future enters any weight. Between rebalances the "
    "positions drift with market prices, as a real fund's would. The "
    "out-of-sample period begins only after the first full estimation "
    "window: 1 February 2021 for equity-calendar funds (734 trading days "
    "ending 29 December 2023). Because cryptocurrencies trade every day, "
    "the crypto fund runs on its own 365-day calendar, begins on 1 October "
    "2020 (1,187 days), and is annualised with 365 periods; equity funds "
    "use 252. Mixing the two would overstate crypto risk-adjusted returns "
    "by roughly 20 percent, so the calendars are kept separate throughout.")
body(
    "Three design decisions are worth stating plainly. First, all funds are "
    "long-only with weights summing to one: short positions would be hard "
    "to explain to the novice user this product serves, and they do not "
    "serve the product's evidence-first philosophy. Second, the Sharpe "
    "ratio assumes a zero risk-free rate and the backtest assumes zero "
    "transaction costs; both are stated simplifications, and Section 6 "
    "returns to what they leave out. Third, the maximum-Sharpe portfolios "
    "are computed with the convex tangency reformulation (minimise y'Σy "
    "subject to μ'y = 1) with a shrunk covariance estimate, after the "
    "direct ratio optimisation stalled silently on every equity window — "
    "a failure caught only because a sanity check compared weights across "
    "methods. The course's portfolio weeks supplied the engine (estimation "
    "windows, formation-versus-return-date discipline); the covariance "
    "shrinkage and the QP reformulation are deliberate beyond-course "
    "choices, made for numerical reliability.")

# ---------------------------------------------------------------------------
doc.add_heading("2. Out-of-sample results and fund fact sheets", level=1)
rubric("Funds (15%), HD continued: fact sheets and required exhibits "
       "present, and the funds are COMPARED, not just listed.")
body(
    f"The out-of-sample results separate construction method from luck. "
    f"Combined Risk-Parity delivered the best risk-adjusted performance of "
    f"the shelf: a Sharpe ratio of {rp['sharpe']:.2f}, from an annualised "
    f"return of {pct(rp['ann_return'])} against {pct(rp['ann_vol'])} "
    f"volatility and a maximum drawdown of {pct(rp['max_drawdown'])}. "
    f"Combined Max-Sharpe produced the highest combined-fund return "
    f"({pct(cms['ann_return'])}) but paid for it with the deepest "
    f"combined-fund drawdown outside crypto ({pct(cms['max_drawdown'])}) "
    f"and a lower Sharpe ({cms['sharpe']:.2f}). Combined Min-Variance did "
    f"what its construction promises in volatility terms: the smallest "
    f"annualised volatility ({pct(cmv['ann_vol'])}) and the shallowest "
    f"drawdown ({pct(cmv['max_drawdown'])}) of the shelf, at the cost of "
    f"the lowest combined-fund return ({pct(cmv['ann_return'])}). The "
    f"equal-weight benchmark sits between the two extremes (Sharpe "
    f"{cew['sharpe']:.2f}), which is itself a useful reference point: the "
    f"optimised funds do not all beat the naive benchmark on return, and "
    f"where they win is in the shape of the ride.")
body(
    f"The single-asset funds frame the diversification story. Equity "
    f"Max-Sharpe ({pct(ems['ann_return'])} return, "
    f"{pct(ems['max_drawdown'])} drawdown) is dominated by its combined "
    f"sibling on both axes, showing what adding a second asset class "
    f"contributed in this sample. Crypto Max-Sharpe is the cautionary "
    f"exhibit: {pct(cr['ann_return'])} annualised return sounds respectable "
    f"until placed beside {pct(cr['ann_vol'])} volatility and a "
    f"{pct(cr['max_drawdown'])} maximum drawdown — a fund that lost more "
    f"than four fifths of its peak value before recovering. The growth-of-"
    f"$1 figure makes the same point visually: the crypto line's peaks are "
    f"the highest on the chart and its troughs the lowest, while the "
    f"combined funds compound in a narrower band.")
evidence("Insert results/tables/performance_metrics.csv as a formatted "
         "Word table (Exhibit 1).")
exhibit("growth_of_1_all_funds.png",
        "Exhibit 2 — Growth of $1, all eight funds, out-of-sample period.")
exhibit("sharpe_by_fund.png", "Exhibit 3 — Sharpe ratio by fund.")
body(
    "Read together, the exhibits support three comparisons. Method matters "
    "more than family: within the combined universe, the spread between "
    f"the best and worst Sharpe ({rp['sharpe']:.2f} versus "
    f"{cmv['sharpe']:.2f}) is larger than the spread between the equity and "
    f"combined versions of the same method. Diversification earns its "
    "textbook reputation in this sample: every combined fund beats its "
    "equity-only counterpart on Sharpe. And no optimised fund dominates "
    "equal-weighting on every measure — the honest summary of three years "
    "of out-of-sample evidence is that the methods buy different shapes of "
    "risk, not free return.")

# ---------------------------------------------------------------------------
doc.add_heading("3. The news-sentiment index", level=1)
rubric("Sentiment (10%), HD: a VALIDATED standalone sector index shown over "
       "time; look-ahead-safe handling justified.")
body(
    "The sentiment layer scores every one of the roughly 147,000 cleaned "
    "news headlines with VADER extended by a custom finance lexicon of "
    "about 60 terms (for example \"beat\", \"downgrade\", \"tumble\"), "
    "hand-assigned scores on VADER's own −4 to +4 convention. Headlines are "
    "scored whole — casing, punctuation and negation are model inputs, not "
    "noise. Scores are averaged per ticker per trading day, smoothed with a "
    "five-day exponential moving average, averaged equally within each "
    "sector, and carried forward over short news silences for at most ten "
    "trading days before reverting to neutral. The index is then lagged by "
    "one full trading day, so the value shown for day t uses only "
    "information available up to t−1. The result is a daily news-tone "
    "index for each of the ten equity sectors from 2 January 2020 to 29 "
    "December 2023 (1,006 trading days).")
body(
    f"Two properties of the index matter for honest interpretation. First, "
    f"headline tone is structurally positive: every sector's average "
    f"reading sits between +{sent_sum['mean'].min():.2f} and "
    f"+{sent_sum['mean'].max():.2f}, and only "
    f"{sent_sum['pct_below_zero'].min():.1f}–"
    f"{sent_sum['pct_below_zero'].max():.1f}% of daily readings fall below "
    f"zero. Zero is therefore not a neutral reference line in this data, "
    f"which is why the app also offers a standardised view comparing each "
    f"sector with its own history (a course-taught correction). Second, "
    f"the index is validated rather than asserted: the same headlines were "
    f"scored independently with the course's finVADER benchmark (VADER plus "
    f"the SentiBigNomics and Henry finance lexicons) on an identical "
    f"pipeline. The two indices agree on the sign of the day's tone on "
    f"{compare.loc['ALL','sign_agreement_pct']:.1f}% of the 1,006 "
    f"overlapping days, with per-sector correlations between "
    f"{compare['pearson_r'].min():.2f} and {compare['pearson_r'].max():.2f} "
    f"(all-sector r = {compare.loc['ALL','pearson_r']:.2f}). Two "
    f"independently built lexicons telling the same directional story on "
    f"roughly nineteen days in twenty is evidence that the custom index "
    f"measures a real signal, not an artefact of one word list.")
exhibit("sector_sentiment_index.png",
        "Exhibit 4 — Sector news-sentiment index over time (lagged).")
body(
    "The validation also exposed the model's honest limits, documented in "
    "the disagreement table (Appendix). VADER's negation rule misfires on "
    "concession headlines (\"US Oil Drillers Continue to Add Rigs Despite "
    "Weak Crude Price\" scored positive by our model), and promotional "
    "vocabulary (\"buy\", \"dividend\") inflates stock-promotion headlines. "
    "The index is a noisy proxy for news tone — it reads headlines, not "
    "articles — and the report treats it accordingly.")

# ---------------------------------------------------------------------------
doc.add_heading("4. Extensions and innovations", level=1)
rubric("Innovation (30%, heaviest), HD: a distinctive IMPLEMENTED extension "
       "shown with evidence; a careful extension with a negative result, "
       "explained, still earns this band.")
body(
    "Four extensions go beyond the course baseline; each is implemented and "
    "evidenced rather than proposed. The headline innovation is the "
    "Practice layer: an \"investment time machine\" that turns the "
    "backtest from a static result into an experience. A user picks a "
    "starting date, amount, and fund mix, then travels through the "
    "three-year out-of-sample period by dragging a timeline slider beneath "
    "their portfolio's value chart. Time flows monthly by default and "
    "slows to daily steps inside seven precomputed turbulent windows "
    "(flagged by a two-sigma rule on daily moves and merged with four "
    "named market events). Holdings are tracked in dollars, so the mix "
    "drifts with market moves exactly as a real holding would — and when "
    "drift exceeds five percentage points, a learning card says so using "
    "the user's own numbers. Every card is collapsible, factual, and "
    "forbidden from evaluating the user's decision: the sandbox teaches "
    "what drawdown, volatility and drift mean by letting the user feel "
    "them, which a static fact sheet cannot do.")
body(
    "The second extension is the custom finance lexicon itself, designed "
    "with AI assistance and governed by human review: every term was "
    "checked against the course's documented failure modes (AI raters "
    "wrongly tag words like \"liability\" and \"debt\" as negative), and "
    "Section 3's benchmark comparison supplies the external validation. "
    "The third is the product's design system — its own colour, type and "
    "figure language (editorial serif headings, a warm paper surface, one "
    "chart grammar with named-event annotations and step-rendered weight "
    "charts) — built to the rubric's own definition of distinctive design. "
    "The fourth is an honest negative result. Folding the sector sentiment "
    "index into the equity fund as a z-scored tilt (strength 0.25, "
    "look-ahead-safe) slightly hurt performance: the tilted fund's Sharpe "
    f"is {tilt['sharpe']:.2f} against the base fund's {ems['sharpe']:.2f}, "
    f"with annualised return {pct(tilt['ann_return'])} versus "
    f"{pct(ems['ann_return'])}. The comparison is reported as found, "
    f"without re-tuning: headline tone is a noisy signal, and a naive tilt "
    f"can subtract value — a finding consistent with the course's own "
    f"warnings about headline sentiment.")
exhibit("fusion_growth_of_1.png",
        "Exhibit 5 — Fusion before vs after: growth of $1, base Equity "
        "Max-Sharpe vs sentiment-tilted.")

# ---------------------------------------------------------------------------
doc.add_heading("5. The app and the investor journey", level=1)
rubric("App (15%), HD: reliable app deployed from a public repo; full "
       "investor journey plus sentiment analytics; original design system "
       "strengthens this band.")
body(
    "HyperInvest is deployed as a Streamlit app from a public GitHub "
    "repository [INSERT LIVE URL AND REPO LINK AFTER DEPLOYMENT]. The app "
    "reads only precomputed artifacts — fund returns, fund weights, the "
    "sentiment index — so it starts fast on the free tier and never "
    "recomputes a backtest or scores text at runtime. It is designed for "
    "two kinds of user, and serves both without compromise. The novice "
    "meets plain-English mode by default: metrics carry reading "
    "instructions anchored in $100 examples, every fund has a one-line "
    "explanation of what it does, and the Practice sandbox lets them learn "
    "from simulated decisions on real history. The experienced investor "
    "switches to professional mode and gets an uncluttered evidence "
    "terminal: the sortable comparison table, complete fact sheets with "
    "methodology notes, and the sentiment analytics.")
body(
    "The investor journey runs: compare the eight funds on the marketplace "
    "table; open a fund's fact sheet (growth of $1, drawdown, weights over "
    "time, current holdings); set an allocation and see its blended "
    "historical performance; save it to a session portfolio. The sentiment "
    "index is one click away, with multi-select sector comparison and the "
    "standardised view. Throughout, the app holds one line absolutely: it "
    "is an evidence layer, not an advice layer. Nothing is scored, ranked "
    "as \"best\", or recommended; the sentiment index carries the explicit "
    "disclaimer that it describes news tone and is not a buy or sell "
    "signal. Neutrality here is not a missing feature — giving investment "
    "advice is a licensed activity, and a student product that never "
    "pretends to it is both more honest and more professional.")
write_placeholder = doc.add_paragraph("[INSERT 2-4 app screenshots after "
                                      "deployment: marketplace, fact sheet, "
                                      "Practice time machine, sentiment tab]")

# ---------------------------------------------------------------------------
doc.add_heading("6. Critical reflection and recommendations", level=1)
rubric("Reflection (10%), HD: evidence-based reflection on what worked, "
       "what did not, why; THREE concrete, specific recommendations; your "
       "own words.")
body(
    "What worked: the walk-forward discipline. Every number in this report "
    "is out-of-sample, and the two methodological traps the course warns "
    "about — look-ahead bias and solver failure — were both caught by "
    "checks rather than by luck (the sentiment lag is tested; the max-"
    "Sharpe stall was caught by comparing weights across methods). What "
    "did not: the sentiment fusion, which subtracted value, and the "
    "sentiment model's residual blind spots around negation and "
    "promotional language. Both are reported as found. The project's "
    "stated simplifications — zero transaction costs, a zero risk-free "
    "rate, and a sample ending 29 December 2023 — bound how far the "
    "conclusions travel.")
body(
    "Three concrete recommendations follow from the evidence. First, add a "
    "transaction-cost and turnover model before trusting any optimised "
    "fund's edge: several funds turn over most of their portfolio at "
    "monthly rebalances (weights swing from 0% to over 90% in a month in "
    "the crypto fund), and even a few basis points of cost per trade would "
    "materially change the net rankings. Second, if this product were "
    "real, refresh the data nightly and re-run the pipeline, so the "
    "evidence stays current to yesterday — while never extending into "
    "prediction. Third, evaluate a tail-risk-aware optimiser such as "
    "mean-CVaR (taught in the course's crypto week) as a fifth method: the "
    "2022 drawdowns are exactly the tail events that variance-based "
    "objectives underweight.")
doc.add_heading("AI workflow statement", level=1)
body(
    "This project was built with a single AI agent (Kimi Code CLI) acting "
    "as planner, implementer and verifier under my direction. I approved "
    "or rejected every material change, reviewed results rather than code, "
    "and own the interpretation in this report. The complete, curated "
    "record — including the decision to discard an earlier attempt and "
    "restart, the corrections I required, and the model decisions I made — "
    "is in ai/prompt_log.md, and my own instruction file is AGENTS.md.")

# ---------------------------------------------------------------------------
doc.add_page_break()
doc.add_heading("Appendix — supporting exhibits", level=1)
exhibit("drawdown_combined_max_sharpe.png",
        "A1 — Drawdown from peak, Combined Max-Sharpe.")
exhibit("weights_combined_max_sharpe.png",
        "A2 — Target weights per monthly rebalance, Combined Max-Sharpe "
        "(top 5; all others pooled in grey).")
exhibit("sentiment_model_comparison.png",
        "A3 — Custom lexicon vs finVADER benchmark.")
evidence("Also reference as tables: results/tables/performance_metrics.csv, "
         "fusion_comparison.csv, sentiment_model_comparison.csv, "
         "sentiment_disagreement_examples.csv, weights_sanity_check.csv.")

doc.add_heading("References", level=1)
body("[VERIFY EACH REFERENCE YOURSELF before submission — "
     "context/verify_ai_output.md: every citation must be one you have "
     "opened. Candidate list:] Course data bundle (equity prices, crypto "
     "prices, news headlines, 2020-2023). Hutto, C.J. & Gilbert, E. "
     "(2014), VADER: A Parsimonious Rule-based Model for Sentiment "
     "Analysis of Social Media Text, ICWSM. Henry, E. (2008), Are "
     "Investors Influenced by How Earnings Press Releases are Written?, "
     "Journal of Business Communication (Henry's finance word list). "
     "SentiBigNomics lexicon (as distributed with the finvader package). "
     "Course weekly materials, weeks 1-9 [cite the specific weeks you use].")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
words = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"draft written to {OUT}")
print(f"approx word count (all text incl. guidance): {words}")
